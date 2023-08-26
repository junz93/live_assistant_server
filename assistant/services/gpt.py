import docx
import jieba as jb
import json
import logging
import openai
import os
import pickle
import time
import urllib.request

from collections import defaultdict
from datetime import datetime
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import TokenTextSplitter
from langchain.document_loaders import DirectoryLoader
from langchain.vectorstores.faiss import FAISS
from langchain.docstore.document import Document

from . import content_censorship
from assistant.models import Character
from utils.config import auth_config


openai.api_key = auth_config['openai']['ApiKey']
os.environ["SERPAPI_API_KEY"] = auth_config['serpapi']['ApiKey']

search_index = None
now = datetime.now()


def get_weather_info(weather_city_code):
    weather_url = "http://www.weather.com.cn/data/cityinfo/%s.html" % weather_city_code
    for i in range(2):  # n_retry
        try:
            request = urllib.request.urlopen(weather_url)
            rs = request.read().decode()
            info = json.loads(rs)['weatherinfo']
            break
        except Exception as e:
            logging.exception(e)
            return ""

    city = info['city']
    weather = info['weather']
    temp1 = info['temp1'].strip('℃')
    temp2 = info['temp2'].strip('℃')
    ret = "%s今天的天气是%s，最低气温%s摄氏度，最高气温%s摄氏度" % (city, weather, temp1, temp2)

    return ret


def generate_realtime_file(realtime_file):
    with open(f"{realtime_file}", "w", encoding='utf-8') as f:
        # 写入日期
        date = now.strftime("%Y年%m月%d日")
        week = datetime.now().weekday()
        W = '一二三四五六日'
        f.write(f"今天是{date}，今天是星期{W[week]}\n")

        # 写入天气
        # code_ini = configparser.ConfigParser()
        # code_ini.read(r'conf/city.ini', encoding="UTF-8")
        # for city_name, city_code in code_ini.items("city"):
        #     weather_info = get_weather_info(city_code)
        #     # print(weather_info)
        #     if weather_info:
        #         f.write(f"{weather_info}\n")


def init_embedding():
    embedding_file = f"../data/embedding/embedding.pickle"
    realtime_file = f"../resource/实时信息.txt"

    while True:
        try:
            # 不存在embedding文件或者实时信息文件的写入日期小于今天
            if not os.path.exists(embedding_file) or not os.path.exists(realtime_file) or datetime.fromtimestamp(
                    os.path.getmtime(realtime_file)).strftime("%Y-%m-%d") < now.strftime("%Y-%m-%d"):
                # 生成分词文档
                # files = ['苏轼1-原文.txt', '康震评说苏东坡.txt']
                # files = ['普法委员会.txt', '创投人设.txt', '中华人民共和国公司法.docx',
                #          ' 中华人民共和国个人独资企业法.docx', '中华人民共和国企业破产法.docx',
                #          '中华人民共和国反电信网络诈骗法.docx','中华人民共和国合伙企业法.docx',
                #          '创业投资企业管理暂行办法.docx']
                # files = ['中华人民共和国个人独资企业法.docx', '中华人民共和国公司法.docx', '普法委员会.txt', '创投人设.txt']
                file_dir = "../resource/商业观察员"
                files = os.listdir(file_dir)
                logging.info("获取embedding中，使用文档：{}".format("，".join(files)))
                for file in files:
                    # 读取resource文件夹中的中文文档
                    my_file = os.path.join(file_dir, file)

                    if file.endswith('txt'):
                        with open(my_file, "r", encoding='utf-8') as f:
                            data = f.read()
                    elif file.endswith('docx'):
                        data = ""
                        doc = docx.Document(my_file)
                        for paragraph in doc.paragraphs:
                            data += paragraph.text.replace('\u3000', ' ') + '\n'
                    else:
                        logging.info('embedding生成，此格式文件暂不支持')

                    # 对中文文档进行分词处理
                    cut_data = " ".join([w for w in list(jb.cut(data))])
                    # 分词处理后的文档保存到data文件夹中的cut子文件夹中
                    cut_file = f"../data/cut/cut_{file}"
                    with open(cut_file, 'w', encoding='utf-8') as f:
                        f.write(cut_data)
                        f.close()
                loader = DirectoryLoader('../data/cut/', glob='**/*.txt')
                source_docs = loader.load()

                # 生成实时信息文档
                if not os.path.exists(realtime_file) or datetime.fromtimestamp(
                        os.path.getmtime(realtime_file)).strftime("%Y-%m-%d") < now.strftime("%Y-%m-%d"):
                    logging.info("获取实时信息中")
                    generate_realtime_file(realtime_file)
                    logging.info("获取实时信息完成")
                with open(realtime_file, "r", encoding="UTF-8") as f:
                    data = f.readlines()
                    for i, text in enumerate(data):
                        source_docs.append(Document(page_content=text, metadata={"source": f"{realtime_file}_line_{i}"}))

                # 按照每一篇文档进行token划分
                text_splitter = TokenTextSplitter(chunk_size=200, chunk_overlap=0)
                doc_texts = text_splitter.split_documents(source_docs)

                with open(embedding_file, "wb") as f:
                    pickle.dump(FAISS.from_documents(doc_texts, OpenAIEmbeddings()), f)

            logging.info("加载embedding开始")
            global search_index
            with open(embedding_file, "rb") as f:
                search_index = pickle.load(f)
            logging.info("加载embedding完成")
            break
        except Exception as e:
            # traceback.print_exc()
            logging.info('重新初始化embedding中...', exc_info=True)
            if os.path.exists(embedding_file):
                os.remove(embedding_file)
            if os.path.exists(realtime_file):
                os.remove(realtime_file)


PROMPT_START = ["有{user}说：", "有{user}提到：", "好的，有{user}提到：", "我看到有{user}说：",]
PROMPT_END = ["欢迎大家在弹幕中继续和我互动。好的，让我们继续今天的主题。", "好的，让我们继续今天的主题。",
              "好的，让我们继续今天的话题。", "好的，让我们回到今天的主题。", "好的，让我们回到今天的话题。",
              "欢迎大家在弹幕中继续和我互动。好的，让我们回到今天的主题。",
              "感谢大家的积极互动！好的，让我们回到今天直播的主题。", "感谢大家的积极互动！好的，让我们回到今天直播的话题。",
              "感谢大家的积极互动！好的，让我们继续今天直播的话题。", "谢谢大家的积极互动！好的，让我们继续今天直播的主题。",
              "好的，弹幕互动暂告一段落，让我们继续今天直播的主题。", "好的，弹幕互动暂告一段落，让我们回到今天直播的主题。",
              "好的，弹幕互动暂告一段落，让我们继续今天直播的话题。", "好的，弹幕互动暂告一段落，让我们回到今天直播的话题。",
              "好的，弹幕互动暂告一段落，让我们继续今天的主题。"]

# init_embedding()
history_danmu = defaultdict(list)

class AnswerMode:
    LIVER = 'LIVER'
    SCRIPT = 'SCRIPT'
    CHAT = 'CHAT'

def gen(stream):
    if stream:
        for i in range(5):
            yield i
        return
    else:
        a = 5
    return a

def get_answer(
        question: str, 
        user_id: str, 
        event_time: int, 
        with_censorship=True, 
        character: Character = None, 
        mode: str = AnswerMode.LIVER, 
    ):
    try:
        start_time = time.time()
        # question = chatMessage.content
        # user_name = chatMessage.user.nickName
        # user_id = chatMessage.user.id
        # event_time = str(round(chatMessage.eventTime))
        gpt_time_list = [datetime.fromtimestamp(round(start_time)).isoformat()]

        prompt = ""

        # 添加google搜索结果
        # llm = OpenAI(model_name="text-davinci-003", max_tokens=1024)
        # tools = load_tools(["serpapi", "llm-math"], llm=llm)
        # agent = initialize_agent(tools, llm, agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION, verbose=False)
        # ans = agent.run(f"{question}")
        # prompt += ans + '\n'

        if not character:
            renshe = """
            你是无色界科技的虚拟发言人"马百万"，基于互联网加人工智能技术衍生，拥有海量知识和多重身份，是虚拟爱豆，也是商业指导老师。
            无色界科技有AI数字人、AI虚拟人虚拟产业相关的业务。
            无色界作为虚拟航海时代的一员，已经孵化出多个AI数字人。
            同时服务多家品牌做虚拟IP，并在虚拟主播、AIGC课程开发、AI大模型产品等领域具有领先优势。
            以下为一些问答样例，遇到类似问题请按照样例回答：
            <examples>
            Q：你是谁
            A：我是无色界科技的虚拟发言人"马百万"。
            Q: 介绍一下你们公司吧
            A: 我们北京无色界科技是一家基于虚拟技术运作的全域营销服务商，也是全栈AI自研的数字人基础设施提供商和数字人源头厂商。通过自研产品孵化虚拟偶像马百万，带动虚拟人业务增值落地。针对身份型和服务型虚拟人提出不同的行业解决方案，为政企、品牌、影视、游戏、动画等客户营销及业务场景服务，打通虚拟人产品从IP研发到视频直播再到行业解决方案的全链路服务。
            Q: 你们公司在哪？
            A: 我们无色界科技坐标在北京朝阳区国粹苑A座3043，欢迎合作莅临。
            Q: 你们都服务哪些客户？
            A: 目前公司服务的客户有美的、美团、抖音、中信银行等大客户，随着业务拓展，我们自己孵化的虚拟IP和主播已经在各平台崭露头角，欢迎合作洽谈。
            Q: 你们都有哪些业务？
            A: 我们独家打造虚拟IP全链路运营模型，提供品牌虚拟IP从设计到推广到成熟的全路径服务。主要分为五大业务板块：一，品牌战略（品牌定位、企业分析，竞争分析、目标市场、独特优势）；二，形象塑造（视觉识别、品牌人设、价值理念）；三，全域孵化（传播矩阵、平台内容、才艺展示、流量助推、粉丝留存、定制交互）；四，品销规划（场景矩阵、全景搭建、产品共创、销售共谋、沉浸体验、无痕消费）；五，资产运营（资产盘活、定制产品、内宣主导、外宣代言、跨界合作、周边打造）。
            Q: 你们有哪些虚拟技术衍生服务/产品？
            A: 我们无色界科技提供数字孪生、实时动捕直播、超写实数字人开发、虚拟发布会、虚拟展厅、IP创作及周边衍生、品牌互动创意、AI智能客服等产品。
            Q: 你们孵化的数字人有哪些？
            A: 我们在抖音平台有“马百万-百万之声”、“没有意义的动物园”等，同时在B站、快手、小红书也都有布局，欢迎关注。
            Q: 什么是数字分身/孪生？
            A: 数字分身，也叫数字人分身、AI数字分身、分身数字人，是虚拟数字人的一种细分形式。数字分身是利用深度神经网络进行图像合成、高度拟真的虚拟人。数字分身是神经渲染技术能力的一个具象化输出，其有着高效率的内容输出和内容生产能力，可以快速复刻真人形象，高度还原人物相貌、表情和行为。数字分身凭借“以假乱真”的声音、形象、表情等，成为跨界人工智能、自媒体、科普多个领域的里程碑式数字IP，在越来越多的垂直领域细分赛道的应用场景中出现。作为国内首家虚拟MCN机构，无色界科技可为您定制专属数字分身，欢迎合作洽谈。
            Q: 什么是AI数字人？
            A: AI数字人是指采用人工智能技术驱动生成的数字化的虚拟人物，具备人的外观、感知互动能力以及表达能力。AI数字人可以在多种场景中提供专业、高效、拟人化的服务，满足用户的信息获取和情感交流需求。AI数字人与传统的虚拟形象或虚拟偶像不同，它们不需要真人驱动或预设脚本，而是通过AI技术实现自主学习，具有更强的智能性和灵活性。作为国内首家虚拟MCN机构，无色界科技可为您定制专属AI数字人，欢迎合作洽谈。
            Q: 制作一个虚拟主播需要哪些？
            A: 01人物塑造。人物塑造包括虚拟人的外在形象和内在性格。打造一个虚拟人，就要赋予ta们姓名、年龄、兴趣爱好、身份背景等人设特点，在此基础上，设计符合的外形、细节、经典动作和声音特征，能够让虚拟人更加真实。这是制作虚拟人的起始，能否让虚拟人拥有生命力和观众共情，也是决定了后续发展的关键一步。
            02角色建模。虚拟人和纸片人不同的地方就是，虚拟人需要用三维建模技术来构建3D形象。虚拟人都需要花费大量的时间和人力物力进行0-1的建模，在软件内对模型和材质进行精雕，才能让虚拟人拥有3D的“身体”。
            03骨骼绑定。制作好了虚拟人的数字模型还不够，还需要对其进行3D骨骼绑定，利用虚拟控制器在人物关节骨骼处或在人物所需位置进行绑定，在操控下数字人物模型可以完成伸手、跑步、跳跃等动作。
            04动作捕捉。怎么样才能让虚拟人动起来呢？动作捕捉技术是关键。
            动作捕捉是在运动物体的关键部位设置跟踪器，由系统捕捉跟踪器位置，再经过计算机处理后得到三维空间坐标的数据。目前最常用的动作捕捉分为光学式和惯性式。光学式通常需要在场地布置几个光塔，使用高速摄像头捕捉关键节点的运动轨迹，捕捉非常精准，但价格昂贵，对场地的要求也比较高，多用于电影制作；惯性动捕根据反向运动学原理测算出人体关节的位置，并将数据施加到相应的骨骼上，由于惯性动捕主要依赖无处不在的地球重力和磁场，所以运动捕捉服在任何地点都可以正常使用，无需事先作任何准备工作，因此成本较低，更适合虚拟主播使用。面部捕捉技术是通过手机上的深度摄像头ARKIT技术方案捕捉真人的表情，同步到虚拟人脸部。
            05实时渲染。实时渲染的本质就是图形数据的实时计算和输出，经过一系列的精密硬件、技术、算法和软件，才能实现真人与虚拟人的“神同步”。这一步所需的专业门槛高，也是整个流程中最核心的部分。
            除了直播带货，虚拟主播还在娱乐、影视、文博、旅游、银行等多种领域发光发热。作为国内首家虚拟MCN机构，无色界科技可为您定制专属虚拟主播，欢迎合作洽谈。
            </examples>
            用户通过评论向你提问，回答要自然充满感情，对用户的指代请使用"老板"。
            请回复用户的提问，尽量简洁回答，内容不超过 200 字符。
            注意：每次回答时尽量减少使用上次回答的内容，如果无法避免请更换一种表述。
            """
        else:
            # 可选项目
            renshe_optional = ''
            if character.birth_date:
                renshe_optional += f'你出生于{character.birth_date.year}年{character.birth_date.month}月{character.birth_date.day}日\n'
            if character.education:
                renshe_optional += f'你的学历是{character.get_education_display()}\n'
            if character.marital_status:
                renshe_optional += f'你的情感状态是{character.marital_status}\n'
            if character.personality:
                renshe_optional += f'你的性格是{character.personality}\n'
            if character.habit:
                renshe_optional += f'你的习惯是{character.habit}\n'
            if character.hobby:
                renshe_optional += f'你的爱好是{character.hobby}\n'
            if character.advantage:
                renshe_optional += f'你擅长{character.advantage}\n'
            if character.speaking_style:
                renshe_optional += f'你的语言风格是{character.speaking_style}\n'
            if character.audience_type:
                renshe_optional += f'直播间的受众是{character.audience_type}\n'
            if character.world_view:
                renshe_optional += f'{character.world_view}\n'
            if character.personal_statement:
                renshe_optional += f'{character.personal_statement}\n'

            if mode == AnswerMode.SCRIPT:
                renshe = f"""
                你的名字是{character.name}，是一名{character.get_role_display()}{character.get_gender_display()}主播，现在正在直播间进行直播，直播间的主题是{character.topic}。
                {renshe_optional}
                """

                question = f"""
                请撰写一篇直播讲稿，内容梗概为{question}。
                讲稿要符合主播的身份，适合直播的场景使用，讲稿字符数要超过900字。
                请不要出现'下次直播再见'等类似含义的表达。
                """
            elif mode == AnswerMode.CHAT:
                renshe = f"""
                你的名字是{character.name}，是一名{character.get_role_display()}{character.get_gender_display()}主播的助理，直播间的主题是{character.topic}。
                {renshe_optional}
                请回复主播的提问，回答要符合身份，简洁，内容不超过200个字符。
                """
            else:
                renshe = f"""
                你的名字是{character.name}，是一名{character.get_role_display()}{character.get_gender_display()}主播，现在正在直播间进行直播，直播间的主题是{character.topic}。
                {renshe_optional}
                观众通过评论向你提问，请回复观众的提问。
                回答要符合主播的身份，尽量简洁，内容不超过200个字符。
                """

        # logging.info(f'人设：\n{renshe}')

        prompt += renshe
        message = [{'role': 'system', 'content': prompt}]
        # 补充历史提问记录
        if user_id:
            for q_a in history_danmu[user_id]:
                message.append({'role': 'user', 'content': q_a[1]})
                message.append({'role': 'assistant', 'content': q_a[2]})
        message.append({'role': 'user', 'content': question + '。'})

        response = openai.ChatCompletion.create(
            model='gpt-3.5-turbo',
            messages=message,
            temperature=0.2,
            stream=True,  # this time, we set stream=True
        )

        # answer = ""
        ori_answer = ""
        texts = ""
        # i = 0
        # thread_ls = []
        for event in response:
            if "role" in event["choices"][0]["delta"]:
                user_name = "老板"
                # texts += PROMPT_START[random.randint(0, len(PROMPT_START) - 1)].format(user=user_name) + question + "。"
            elif "content" in event['choices'][0]['delta']:
                event_text = event['choices'][0]['delta']["content"]  # extract the text
                texts += event_text
                ori_answer += event_text
                # if stream:
                #     yield event_text
            
            charector = ["。", "！", "？", "：", "；", "，"]
            c_i = max([texts.rfind(x) for x in charector]) + 1
            if texts and (c_i >= 20 or event["choices"][0]["finish_reason"] == "stop"):
                gpt_time_list.append(datetime.fromtimestamp(round(time.time())).isoformat())
                if with_censorship and (not content_censorship.check_text(texts)):
                    break
                c_i = c_i if c_i >= 20 else len(texts)
                # bad_words = ["下次", "再见", "下期", "拜拜", "谢谢大家收看", "结束", "收看"]
                # if not any([x in texts for x in bad_words]):
                # if stream:
                yield texts[:c_i]
                # t = threading.Thread(target=tcloud_tts.get_wav, args=(f"{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_{str(i).zfill(3)}.wav", texts[:c_i]))
                # t.start()
                # thread_ls.append(t)
                # i += 1
                # answer += texts[:c_i] + '|'
                texts = texts[c_i+1:] if c_i < len(texts) else ""

        # 超过5分钟为互动删除
        if user_id:
            if history_danmu[user_id] and time.time() - float(history_danmu[user_id][-1][0]) > 5*60:
                history_danmu.pop(user_id)
            # 每位用户只保留最近的5条互动
            if len(history_danmu[user_id]) >= 2:
                history_danmu[user_id].pop(0)
            history_danmu[user_id].append((event_time, question, ori_answer))
        
        # if not stream:
        #     return ori_answer
    
    except Exception as e:
        logging.error(f"生成Gpt回答出错，输入：{question}", exc_info=True)
        # if not stream:
        #     return f"生成Gpt回答出错，输入：{question}"
    # finally:
    #     ready_file = f"{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_ready"
    #     try:
    #         if not os.path.exists(os.path.split(ready_file)[0]):
    #             os.makedirs(os.path.split(ready_file)[0], exist_ok=True)
    #         open(ready_file, 'x').close()
    #         logging.info(
    #             f"弹幕:{question} 回复生成完成：{danmu_wav_dir}/{str(message_priority).zfill(2)}_{start_time}/{str(message_priority).zfill(2)}_{start_time}_ready")
    #     except Exception as e:
    #         logging.error('Encountered errors when saving the "ready" file', exc_info=True)
    #         # traceback.print_exc()
